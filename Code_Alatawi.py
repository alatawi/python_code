"""
By Abdulrahman Alatawi - 11/07/2017 - to.alatawi@gmail.com
Using Python 3.6.2
This code is for Expanding Abbreviations in Source Code using
Bayesian (Unigram-Bigram) Based Inference
Note: You need to install the following:
pip3 install wordsegment
pip3 install python-docx
"""
#======== Import ==========================
import os
import glob
import re
import csv
import math 
import itertools
from docx import Document
from docx.shared import Inches, Mm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.text.run import Font, Run
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import defaultdict
from time import gmtime, strftime
from wordsegment import load, segment
load()
#======== Constants here ==================            
inputFolder = "bnl_Method"
outputFolder = "output"
abbsFileName = "abbs.csv"
ngramsDatasets = ['NL_Unigrams.csv', 'SE_Unigrams.csv','NL_Bigrams.csv','SE_Bigrams.csv']
#======== Properties Functions ============
'Get abbreviation type, either CAS OR NAS'
def getAbbType(abb, candidate):
    abbType = 'NAS'
    if abb[0:] == candidate[0:len(abb)]:
            abbType = 'CAS'
    return abbType

'Get probability of abbreviation type, either 0.16 for NAS or 0.84 for CAS'
def get_P_abbType(abbType):
    P_abbType = 0.16
    if abbType == 'CAS':
        P_abbType = 0.84
    return P_abbType

'Get TES ratio'
def getTES(abb, candidate):
    return (1-(len(abb)/len(candidate)))

'Get probability of TES'
def getP_TES(TES):
    P_TES = 0.0000000001
    if TES<=0.09:
        P_TES = 0.008594958
    elif TES <= 0.19:
        P_TES = 0.050616912
    elif TES <= 0.29:
        P_TES = 0.036402174
    elif TES <= 0.39:
        P_TES = 0.059484108
    elif TES <= 0.49:
        P_TES = 0.053300405
    elif TES <= 0.59:
        P_TES = 0.142312666
    elif TES <= 0.69:
        P_TES = 0.188466811
    elif TES <= 0.79:
        P_TES = 0.201689823
    elif TES <= 0.89:
        P_TES = 0.219249205
    elif TES <= 0.99:
        P_TES = 0.039882937
    return P_TES

'Get candidate distance'
def getCandidateDist(abbLineNumber, candidateLineNum):
    return abs(int(abbLineNumber)-int(candidateLineNum))

'Get probability of distance = 0.83*exp(-2.53*C_Dist)'
def getP_Dist(candidateDist):
    return 0.83*math.exp(-2.53*candidateDist)

'Get probability of P(abb|andidate) = P_abbType*P_TES*P_Dist'
def getP_Abb_C(P_abbType, P_TES, P_Dist):
    P_Abb_C = P_abbType*P_TES*P_Dist
    return P_Abb_C

'Get candidate frequency using csv.reader'
def getC_Freq(candidate, datasetName):
    C_Freq = 1 #if not found
    with open(os.getcwd()+"/required_data/"+datasetName) as f:
        dataset = csv.reader(f, delimiter=',')  
        for line in dataset:
            if line[0] == candidate:
                C_Freq = line[1]
                return C_Freq
    return C_Freq

'Get probability of candidate'
def getP_C(C_Freq, totalGrams):
    P_candidate = int(C_Freq) / int(totalGrams)
    return P_candidate
#======== Functions =======================
def makeOutputDir(folderName):
    if not os.path.exists(folderName):
        os.makedirs(folderName)
        
'Get abbs and line numbers with the correct expansion, using csv.reader'
def getAbbsDataDict(path):
    abbsDataDict = {}
    with open(path) as f:
        lines = csv.reader(f, delimiter=',')
        i = 1
        for line in lines:
            innerList = []
            innerList.append(line[0])
            innerList.append(line[1])
            innerList.append(line[2])
            abbsDataDict[i] = innerList
            i = i + 1
    return abbsDataDict

'Get Abbreviation’s N-grams (AN)'
def getAbbANs(abb):
    abbANs = []
    chars = list(abb)
    for t in itertools.product(range(len('01')), repeat=len(chars)-1):
        abbANs.append(''.join([chars[i]+t[i]*'|' for i in range(len(t))])+chars[-1])
    return abbANs

'Get the Abbreviation N-gram Unique Part(ANUP)'
def getANUPs(abbANs):
    ANUPs = []
    abbANs = [words for segments in abbANs for words in segments.split('|')]
    for abbAN in abbANs:
        if abbAN not in ANUPs:
            ANUPs.append(abbAN)
    return ANUPs

'Read source code into list'
def readSourceCodeIntoList(inPath, i):
    fileCorrectName = 'method_'
    if i <= 9:
        fileCorrectName = fileCorrectName + '000'
    elif i <= 99:
        fileCorrectName = fileCorrectName + '00'
    elif i <= 999:
        fileCorrectName = fileCorrectName + '0'
    fileCorrectName = fileCorrectName + str(i) + '.java'
    with open(inPath+fileCorrectName) as f:
        listOfLines = [word.strip() for word in f]
    sourceCodeLines = []
    for line in listOfLines:
        sourceCodeLines.append(line)
    return sourceCodeLines

'Get unigram Candidates-Pairs (CPs): ANUPs with their Candidates and line numbers'
def getCPs(sourceCodeLines, abb, abbLineNumber, ANUPs):#Changed 2/6/2017
    candidatesWithLineNum = {}
    lineNumber = 0
    for line in sourceCodeLines:
        line = ' '.join(re.findall("[a-zA-Z]+", line))
        unwantedWords = ["abstract", "assert", "break", "case", "catch", "class", 
                         "continue", "default", "do", "else", "extends", "false",
                         "final", "finally", "for", "goto", "if", "implements",
                         "import", "instanceof", "interface", "native", "const",
                         "new", "null", "package", "private", "protected", "public",
                         "return", "short", "static", "strictfp", "super", "switch",
                         "synchronized", "this", "throw", "throws", "true",
                         "try", "void", "while"]
        words = line.split()
        filteredWords  = [word for word in words if word.lower() not in unwantedWords]
        strippedLine = ' '.join(filteredWords)
        lineNumber += 1
        listOfPotentialWords = []
        lineWords = strippedLine.split()
        for word in lineWords:
            if (len(word) > len(abb)):
                segmentedWords = segment(word)
                for candidate in segmentedWords:
                    if candidate != abb and (candidate[0] in list(abb)):
                        if candidate in candidatesWithLineNum:
                            newDist = getCandidateDist(abbLineNumber, lineNumber)
                            if newDist < candidatesWithLineNum[candidate]:
                                candidatesWithLineNum[candidate] = newDist
                        else:
                            candidatesWithLineNum[candidate] = getCandidateDist(abbLineNumber, lineNumber)
    CP = {}
    for ANUP in ANUPs:
        currentANUPDict = {}
        for candidate in candidatesWithLineNum:
            if ANUP[0] == candidate[0] and len(candidate[0])>1:
                if len(ANUP) == 1:# Unigram
                    currentANUPDict[candidate] = candidatesWithLineNum[candidate]
                else: #Multigram
                    okayToAdd = 1
                    for ch in list(ANUP):
                        if ch not in list(candidate):
                            okayToAdd = 0
                    if okayToAdd:
                        currentANUPDict[candidate] = candidatesWithLineNum[candidate]
        if currentANUPDict != {}:
            CP[ANUP] = currentANUPDict
    return CP

def getTotalGrams(datasetName):
    totalGrams = 0
    with open(os.getcwd()+"/required_data/"+datasetName) as f:
        lines = csv.reader(f, delimiter=',')
        for line in lines:
            totalGrams +=(int(line[1]))
    return totalGrams

'Computing unigram Candidate-Pair’s Properties (CPP)'
def getCPP(abb, abbLineNumber, datasetName, CPs):
    CPP_Dict = {}
    totalGrams = getTotalGrams(datasetName)
    for ANUP in CPs:
        candidate_Dict = {}
        for candidate in CPs[ANUP]:
            valuesDict = {}
            valuesDict['abbType'] = getAbbType(ANUP, candidate)
            valuesDict['P_abbType'] = get_P_abbType(valuesDict['abbType'])
            valuesDict['TES'] = getTES(ANUP, candidate)
            valuesDict['P_TES'] = getP_TES(valuesDict['TES'])
            valuesDict['candidateDist'] = CPs[ANUP][candidate]
            valuesDict['P_Dist'] = getP_Dist(valuesDict['candidateDist'])
            valuesDict['P_Abb_C'] = getP_Abb_C(valuesDict['P_abbType'], valuesDict['P_TES'], valuesDict['P_Dist'])
            valuesDict['C_Freq'] = getC_Freq(candidate, datasetName)
            valuesDict['P_C'] = getP_C(valuesDict['C_Freq'], totalGrams)
            candidate_Dict[candidate] = valuesDict
        CPP_Dict[ANUP] = candidate_Dict
    return CPP_Dict

'Format candidates with "|"'
def formatCandidates(oneCombination):
    formattedCandidates = oneCombination[0]
    if(len(oneCombination) > 1):
        for c in oneCombination[1:]:
            formattedCandidates = formattedCandidates + '|' + c
    return formattedCandidates

'Get all combinations'
def getAllCombinations(currentCandidatesDict, ngramSize):
    allCombinationsDict = {}
    j = 1
    for c in currentCandidatesDict[1]:
        oneCombination = []
        oneCombination.append(c)
        if(ngramSize > 1):
            for c in currentCandidatesDict[2]:
                oneCombination.append(c)
                
                if(ngramSize > 2):
                    for c in currentCandidatesDict[3]:
                        oneCombination.append(c)
                        
                        if(ngramSize > 3):
                            for c in currentCandidatesDict[4]:
                                oneCombination.append(c)
                                
                                if(ngramSize > 4):
                                    for c in currentCandidatesDict[5]:
                                        oneCombination.append(c)

                                        if(ngramSize > 5):
                                            for c in currentCandidatesDict[6]:
                                                oneCombination.append(c)
                                                allCombinationsDict[j] = formatCandidates(oneCombination)
                                                j = j + 1
                                                del oneCombination[len(oneCombination)-1]
                                        else:
                                            allCombinationsDict[j] = formatCandidates(oneCombination)
                                            j = j + 1
                                        del oneCombination[len(oneCombination)-1]
                                else:
                                    allCombinationsDict[j] = formatCandidates(oneCombination)
                                    j = j + 1
                                del oneCombination[len(oneCombination)-1]
                        else:
                            allCombinationsDict[j] = formatCandidates(oneCombination)
                            j = j + 1
                        del oneCombination[len(oneCombination)-1]
                        
                else:
                    allCombinationsDict[j] = formatCandidates(oneCombination)
                    j = j + 1
                del oneCombination[len(oneCombination)-1]
        else:
            allCombinationsDict[j] = c
            j = j + 1
    return allCombinationsDict

'Get the best candidate dictionary using bigram datasets'
def getANs_Combinations(abb, abbANs, CPs):
    ANs_Combinations = {}
    for AN in abbANs:
        currentCandidatesDict = {}
        AN_Parts = AN.split('|')
        ngramSize = len(AN_Parts)
        i = 1
        okayToAdd = 1
        for AN_Part in AN_Parts:
            if AN_Part not in CPs:
                okayToAdd = 0
            else:
                currentCandidatesDict[i] = CPs[AN_Part]
                i = i + 1
        if currentCandidatesDict != {} and okayToAdd:
            allCombinationsDict = getAllCombinations(currentCandidatesDict, ngramSize)
            combinationsList = []
            for oneCombination in allCombinationsDict:
                combinationsList.append(allCombinationsDict[oneCombination])
            ANs_Combinations[AN] = combinationsList
    return ANs_Combinations

'Get P_Abb_C_Avg'
def get_P_C_Abb_Ave(AN, oneCombination, CPP):
    candidates = oneCombination.split('|')
    AN_Parts = AN.split('|')
    P_C_Abb_List = []
    i = 0
    for AN_Part in AN_Parts:
        P_C_Abb = (float(CPP[AN_Part][candidates[i]]['P_Abb_C']))*(float(CPP[AN_Part][candidates[i]]['P_C']))
        P_C_Abb_List.append(P_C_Abb)
        i +=1
    P_C_Abb_Avg = 1
    for P_C_Abb in P_C_Abb_List:
        P_C_Abb_Avg *= P_C_Abb
    n = len(candidates)
    P_C_Abb_Avg = math.pow(P_C_Abb_Avg, (1/n))
    return P_C_Abb_Avg

'Get Bi_P_Abb_C_Avg' #----------------------
def get_Bi_P_C_Abb_Ave(datasetName, AN, oneCombination, CPP):
    candidates = oneCombination.split('|')
    n = len(candidates)
    AN_Parts = AN.split('|')
    Bi_P_C_Abb_Avg = 0
    if n == 1:
        Bi_P_C_Abb_Avg = (float(CPP[AN_Parts[0]][candidates[0]]['P_Abb_C']))*(float(CPP[AN_Parts[0]][candidates[0]]['P_C']))
    else:
        P_C = 1
        P_Abb_C = 1
        totalGrams = getTotalGrams(datasetName)
        i = 0
        while i < (n-1):
            C_Freq = getC_Freq(candidates[i]+' '+candidates[i+1], datasetName)
            P_C *= getP_C(C_Freq, totalGrams)
            P_Abb_C *= math.pow(float(CPP[AN_Parts[i]][candidates[i]]['P_Abb_C'])*float(CPP[AN_Parts[i+1]][candidates[i+1]]['P_Abb_C']), (1/(n-1)))
            i +=1
        Bi_P_C_Avg = math.pow(P_C, (1/(n-1)))
        Bi_P_Abb_C = math.pow(P_Abb_C, (1/(n-1)))
        Bi_P_C_Abb_Avg = Bi_P_Abb_C * Bi_P_C_Avg
    return Bi_P_C_Abb_Avg

'Get final calulations for all combinations'
def getFinalCalculations(CPP, ANs_Combinations):
    finalCalculationsDict = {}
    i = 1
    for AN in ANs_Combinations:
        for oneCombination in ANs_Combinations[AN]:
            print('Unigram - Calculating for: '+str(AN)+' -> '+str(oneCombination))
            innerList = []
            innerList.append(AN)
            innerList.append(oneCombination)
            innerList.append(get_P_C_Abb_Ave(AN, oneCombination, CPP))
            finalCalculationsDict[i] = innerList
            i += 1
    finalCalculationsDictSorted = sorted(finalCalculationsDict.items(), key=lambda e: e[1][2], reverse=True)
    return finalCalculationsDictSorted

def printResults(document, datasetType, originalWord, CPP, finalCalculationsDict, biFinalCalculationsDict):
    p = document.add_heading('Using ', level=1)
    p.add_run(datasetType+' Dataset').bold = True
    p.add_run(': ')
    document.add_paragraph('Computing unigram Candidate-Pair’s Properties (CPP):', style='List Bullet')
    table = document.add_table(rows=1, cols=11)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ANUP'
    hdr_cells[1].text = 'Candidate'
    hdr_cells[2].text = 'ANUP Type'
    hdr_cells[3].text = 'P(ANUP Type)'
    hdr_cells[4].text = 'TES(ANUP, C)'
    hdr_cells[5].text = 'P(TES)'
    hdr_cells[6].text = 'Dist'
    hdr_cells[7].text = 'P(Dist)'
    hdr_cells[8].text = 'P(ANUP|C)'
    hdr_cells[9].text = 'Frequency'
    hdr_cells[10].text = 'P(C)'
    for ANUP in CPP:
        for candidate in CPP[ANUP]:
            row_cells = table.add_row().cells
            row_cells[0].text = '{:<6}'.format(str(ANUP))
            row_cells[1].text = '{:<20}'.format(str(candidate))
            row_cells[2].text = '{:<10}'.format(str(CPP[ANUP][candidate]['abbType']))
            row_cells[3].text = '{:<15}'.format("{:.10f}".format(float(CPP[ANUP][candidate]['P_abbType'])))
            row_cells[4].text = '{:<15}'.format("{:.10f}".format(float(CPP[ANUP][candidate]['TES'])))
            row_cells[5].text = '{:<10}'.format("{:.5f}".format(float(CPP[ANUP][candidate]['P_TES'])))
            row_cells[6].text = '{:<5}'.format(str(CPP[ANUP][candidate]['candidateDist']))
            row_cells[7].text = '{:<15}'.format("{:.10f}".format(float(CPP[ANUP][candidate]['P_Dist'])))
            row_cells[8].text = '{:<15}'.format("{:.10f}".format(float(CPP[ANUP][candidate]['P_Abb_C'])))
            row_cells[9].text = '{:<15}'.format(str(CPP[ANUP][candidate]['C_Freq']))
            row_cells[10].text = '{:<15}'.format("{:.10f}".format(float(CPP[ANUP][candidate]['P_C'])))
    document.add_paragraph('')
    document.add_paragraph('Finding the best candidate using Unigrams:', style='List Bullet')
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Abb'
    hdr_cells[2].text = 'Candidate'
    hdr_cells[3].text = 'P(Abb|C)*P(C)'
    i = 1
    for combination in finalCalculationsDict:
        row_cells = table.add_row().cells
        row_cells[0].text = '{:<5}'.format(str(i))
        row_cells[1].text = '{:<35}'.format(str(combination[1][0]))
        run = row_cells[2].paragraphs[0].add_run('{:<35}'.format(str(combination[1][1])))
        font = run.font
        if i == 1:
            font.color.rgb = RGBColor(255, 0, 0)#Red
        if str(combination[1][1]) == str(originalWord):
            font.color.rgb = RGBColor(0, 160, 0)#Green
        row_cells[3].text = '{:<35}'.format("{:.20f}".format(float(str(combination[1][2]))))
        i += 1
    document.add_paragraph('Finding the best candidate using Bigrams:', style='List Bullet')
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Abb'
    hdr_cells[2].text = 'Candidate'
    hdr_cells[3].text = 'P(Abb|C)*P(C)'
    i = 1
    for combination in biFinalCalculationsDict:
        #'{:<5} {:<35} {:<35} {:<35}'.format('#', 'Abb', 'Candidate', 'P(Abb|C)*P(C)'
        row_cells = table.add_row().cells
        row_cells[0].text = '{:<5}'.format(str(i))
        row_cells[1].text = '{:<35}'.format(str(combination[1][0]))
        run = row_cells[2].paragraphs[0].add_run('{:<35}'.format(str(combination[1][1])))
        font = run.font
        if i == 1:
            font.color.rgb = RGBColor(255, 0, 0)#Red
        if str(combination[1][1]) == str(originalWord):
            font.color.rgb = RGBColor(0, 160, 0)#Green
        row_cells[3].text = '{:<35}'.format("{:.20f}".format(float(str(combination[1][2]))))
        i += 1
    return 0

def isBiagram(index, line, candidates):
    okToAdd = 1
    n = len(candidates)
    i = 0
    for word in line[index:index+n]:
        if candidates[i] != word:
            okToAdd = 0
        i +=1
    return okToAdd

'Get all bigram combinations'
def getBigramCombinations(abb, sourceCodeLines, ANs_Combinations):
    lines_Dict = {}
    i = 1
    for line in sourceCodeLines:
        line = ' '.join(re.findall("[a-zA-Z]+", line))
        unwantedWords = ["abstract", "assert", "break", "case", "catch", "class", 
                         "continue", "default", "do", "else", "extends", "false",
                         "final", "finally", "for", "goto", "if", "implements",
                         "import", "instanceof", "interface", "native", "const",
                         "new", "null", "package", "private", "protected", "public",
                         "return", "short", "static", "strictfp", "super", "switch",
                         "synchronized", "this", "throw", "throws", "true",
                         "try", "void", "while"]
        words = line.split()
        filteredWords  = [word for word in words if word.lower() not in unwantedWords]
        strippedLine = ' '.join(filteredWords)
        line = ' '.join(segment(strippedLine))
        words = line.split()
        lines_Dict[i] = words
        i +=1
    biAbbANs_Dict = {}
    okayToAdd = 0
    for AN in ANs_Combinations:
        combinations = []
        for oneCombination in ANs_Combinations[AN]:
            if '|' in oneCombination:
                candidates = oneCombination.split('|')
            else:
                candidates = []
                candidates.append(oneCombination)
            for line in lines_Dict:
                if str(candidates[0]) in lines_Dict[line]:
                    index = lines_Dict[line].index(str(candidates[0]))
                    if isBiagram(index, lines_Dict[line], candidates):
                        if oneCombination not in combinations:
                            okayToAdd = 1
                            combinations.append(oneCombination)
        if okayToAdd:#Added this 11/13/2017
            biAbbANs_Dict[AN] = combinations
    if biAbbANs_Dict == {} and abb in ANs_Combinations:
        for candidate in ANs_Combinations[abb]:
            if abb == candidate[:len(abb)]:
                biAbbANs_Dict[abb] = candidate
    return biAbbANs_Dict

'Get final bigram calulations for all combinations'
def getBiFinalCalculations(datasetName, CPP, biAbbANs_Dict):
    finalBiCalculationsDict = {}
    i = 1
    for AN in biAbbANs_Dict:
        for oneCombination in biAbbANs_Dict[AN]:
            print('Bigram - Calculating for: '+str(AN)+' -> '+str(oneCombination))
            innerList = []
            innerList.append(AN)
            innerList.append(oneCombination)
            innerList.append(get_Bi_P_C_Abb_Ave(datasetName, AN, oneCombination, CPP))
            finalBiCalculationsDict[i] = innerList
            i += 1
    finalBiCalculationsDictSorted = sorted(finalBiCalculationsDict.items(), key=lambda e: e[1][2], reverse=True)
    return finalBiCalculationsDictSorted
#======== Main Code =======================
def do(inputFolder, outputFolder, abbsFileName, testName, ngramsDatasets):
    import time
    start_time = time.time()
    print('Runing test: '+str(testName))
    inPath = str(os.getcwd()+'/'+inputFolder+'/')
    outPath = str(os.getcwd()+'/'+outputFolder+'/')
    abbsDataDict = getAbbsDataDict(inPath+abbsFileName)
    timeAndDate = strftime("%m/%d/%Y at %H:%M:%S", gmtime())
    document = Document()
    document.add_heading('Result for "'+testName+'"\n'+timeAndDate, 0)
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = Mm(240)
    section.page_width = Mm(327)
    PCEA_Dict = {}
    PCEA_Final = {'NLU':0, 'SEU':0, 'NLB':0, 'SEB':0}
    for i in abbsDataDict:
        abb = abbsDataDict[i][0]
        print("["+str(i)+"] abb: "+str(abb))
        abbLineNumber = abbsDataDict[i][1]
        originalWord = abbsDataDict[i][2]
        abbANs = getAbbANs(abb)
        print('abbANs: '+str(abbANs))
        ANUPs = getANUPs(abbANs)
        print('ANUPs: '+str(ANUPs))
        sourceCodeLines = readSourceCodeIntoList(inPath, i)
        print('sourceCodeLines: '+str(sourceCodeLines))
        CPs = getCPs(sourceCodeLines, abb, abbLineNumber, ANUPs)
        print('CPs: '+str(CPs))
        NLD_CPP = getCPP(abb, abbLineNumber, ngramsDatasets[0], CPs)
        print('NLD_CPP: '+str(NLD_CPP))
        SED_CPP = getCPP(abb, abbLineNumber, ngramsDatasets[1], CPs)
        print('SED_CPP: '+str(SED_CPP))
        ANs_Combinations = getANs_Combinations(abb, abbANs, CPs)
        print('ANs_Combinations: '+str(ANs_Combinations))
        NLD_finalCalculationsDict = getFinalCalculations(NLD_CPP, ANs_Combinations)
        #print('NLD_finalCalculationsDict: '+str(NLD_finalCalculationsDict))
        SED_finalCalculationsDict = getFinalCalculations(SED_CPP, ANs_Combinations)
        #print('SED_finalCalculationsDict: '+str(SED_finalCalculationsDict))
        bigramCombinations = getBigramCombinations(abb, sourceCodeLines, ANs_Combinations)
        #print('bigramCombinations: '+str(bigramCombinations))
        NLD_biFinalCalculationsDict = getBiFinalCalculations(ngramsDatasets[2], NLD_CPP, bigramCombinations)
        #print('NLD_biFinalCalculationsDict: '+str(NLD_biFinalCalculationsDict))
        SED_biFinalCalculationsDict = getBiFinalCalculations(ngramsDatasets[3], SED_CPP, bigramCombinations)
        #print('SED_biFinalCalculationsDict: '+str(SED_biFinalCalculationsDict))
        if printDetails:
            document.add_heading('Abb: ' + str(abb), level=1)
            document.add_paragraph('Abbreviation’s n-grams: ' + str(abbANs), style='List Bullet')
            document.add_paragraph('Avialable candidates with line numbers are: ' + str(CPs), style='List Bullet')
            printResults(document, 'NL', originalWord, NLD_CPP, NLD_finalCalculationsDict, NLD_biFinalCalculationsDict)
            printResults(document, 'SE', originalWord, SED_CPP, SED_finalCalculationsDict, SED_biFinalCalculationsDict)
        #Print final compareson
        innerDict = {}
        innerDict['Original Word'] = originalWord
        try:
            innerDict['NLU'] = NLD_finalCalculationsDict[0][1][1]
        except IndexError:
            innerDict['NLU'] = 'NULL'
        try:
            innerDict['SEU'] = SED_finalCalculationsDict[0][1][1]
        except IndexError:
            innerDict['SEU'] = 'NULL'
        try:
            innerDict['NLB'] = NLD_biFinalCalculationsDict[0][1][1]
        except IndexError:
            innerDict['NLB'] = 'NULL'
        try:
            innerDict['SEB'] = SED_biFinalCalculationsDict[0][1][1]
        except IndexError:
            innerDict['SEB'] = 'NULL'
        PCEA_Dict[abb] = innerDict
        print("["+str(i)+"] abb: "+str(abb)+", ['NLU'.'"+innerDict['NLU']+"'], ['SEU'.'"+innerDict['SEU']+"'], ['NLB'.'"+innerDict['NLB']+"'], ['SEB'.'"+innerDict['SEB']+"']")
        print("\n")
    'End of for - abb'
    #print('PCEA_Dict: '+str(PCEA_Dict))
    
    'Start PCEA Printing'
    document.add_heading('Compute the Percentage of Correctly Expanded Abbreviations (PCEA):', level=1)
    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Abb'
    hdr_cells[2].text = 'Original Word'
    hdr_cells[3].text = 'NLU'
    hdr_cells[4].text = 'SEU'
    hdr_cells[5].text = 'NLB'
    hdr_cells[6].text = 'SEB'
    i = 0
    for abb in PCEA_Dict:
        i +=1
        row_cells = table.add_row().cells
        row_cells[0].text = '{:<2}'.format(str(i))
        row_cells[1].text = '{:<10}'.format(abb)
        row_cells[2].text = '{:<20}'.format(PCEA_Dict[abb]['Original Word'])
        if PCEA_Dict[abb]['NLU'] == PCEA_Dict[abb]['Original Word']:
            row_cells[3].paragraphs[0].add_run(PCEA_Dict[abb]['NLU']).font.color.rgb = RGBColor(0, 160, 0)#Green
            PCEA_Final['NLU'] +=1
        else:
            row_cells[3].paragraphs[0].add_run(PCEA_Dict[abb]['NLU']).font.color.rgb = RGBColor(255, 0, 0)#Red
        if PCEA_Dict[abb]['SEU'] == PCEA_Dict[abb]['Original Word']:
            row_cells[4].paragraphs[0].add_run(PCEA_Dict[abb]['SEU']).font.color.rgb = RGBColor(0, 160, 0)#Green
            PCEA_Final['SEU'] +=1
        else:
            row_cells[4].paragraphs[0].add_run(PCEA_Dict[abb]['SEU']).font.color.rgb = RGBColor(255, 0, 0)#Red
        if PCEA_Dict[abb]['NLB'] == PCEA_Dict[abb]['Original Word']:
            row_cells[5].paragraphs[0].add_run(PCEA_Dict[abb]['NLB']).font.color.rgb = RGBColor(0, 160, 0)#Green
            PCEA_Final['NLB'] +=1
        else:
            row_cells[5].paragraphs[0].add_run(PCEA_Dict[abb]['NLB']).font.color.rgb = RGBColor(255, 0, 0)#Red
        if PCEA_Dict[abb]['SEB'] == PCEA_Dict[abb]['Original Word']:
            row_cells[6].paragraphs[0].add_run(PCEA_Dict[abb]['SEB']).font.color.rgb = RGBColor(0, 160, 0)#Green
            PCEA_Final['SEB'] +=1
        else:
            row_cells[6].paragraphs[0].add_run(PCEA_Dict[abb]['SEB']).font.color.rgb = RGBColor(255, 0, 0)#Red
    'End PCEA Printing'
    PCEA_P = {'NLU':int(PCEA_Final['NLU']/len(abbsDataDict)*100), 'SEU':int(PCEA_Final['SEU']/len(abbsDataDict)*100), 'NLB':int(PCEA_Final['NLB']/len(abbsDataDict)*100), 'SEB':int(PCEA_Final['SEB']/len(abbsDataDict)*100)}
    print(PCEA_P)
    document.add_paragraph("")
    document.add_paragraph("NLU: "+'{:<2}'.format(str(PCEA_P['NLU']))+"%, SEU: "+'{:<2}'.format(str(PCEA_P['SEU']))+"%, NLB: "+'{:<2}'.format(str(PCEA_P['NLB']))+"%, SEB: "+'{:<2}'.format(str(PCEA_P['SEB']))+"%.")
    finishTime = '{:<3}'.format("{:.2f}".format(float((time.time() - start_time)/60)))
    print("Total run time: %s Minutes" % (finishTime))
    document.add_paragraph("")
    document.add_paragraph("Total run time: %s Minutes" % (finishTime))
    latestSaved = 1
    makeOutputDir(outputFolder)
    if glob.glob(outPath+'*.docx'):
        latestSaved = int(max(glob.glob(outPath+'*.docx'), key=os.path.getctime)[-7:-5])+1
    document.save(outPath+testName+'_Result_Report_'+'{0:02}'.format(latestSaved)+'.docx')
    print('\nDone. Check "'+testName+'_Result_Report_'+'{0:02}'.format(latestSaved)+'.docx'+'" in the output folder!')
    return 0

#======== Main Code =======================
printDetails = 0
testName = str(inputFolder)+' Testing'

do(inputFolder, outputFolder, abbsFileName, testName, ngramsDatasets)
